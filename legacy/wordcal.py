# -*- coding: utf-8 -*-
"""
@Time ： 2026/2/3 10:40
@Auth ： guoshuai
@File ：wordcal.py
@Describe:
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE


def create_formatted_document():
    # 创建文档
    doc = Document()

    # 添加主标题
    title = doc.add_paragraph('XX市实验中等专业学校教学事故认定与处理办法')
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # 设置标题样式
    title_run = title.runs[0]
    title_run.bold = True
    title_run.font.size = Pt(22)
    title_run.font.name = '黑体'

    # 添加空行
    doc.add_paragraph()

    # 定义文本内容
    content = [
        ('第一条', False, True),
        ('为规范教学管理，严肃教学纪律，维护正常教学秩序，增强教师责任意识，保障教学质量，制定本办法。', False, False),
        ('第二条', False, True),
        (
        '教学事故是指教师、教学管理及教学保障人员因责任心不强或工作失误，干扰或损害正常教学秩序的事件。教学事故分为一般教学事故和严重教学事故。',
        False, False),
        ('第三条 凡有下列情形之一者，属于一般教学事故：', False, True),
        (
        '1.  上课、自习辅导、监考等因不可抗力（如恶劣天气、突发事故等）以外的原因迟到、中途离开或提前下课（5分钟以内），且未提前报备的；',
        True, False),
        ('2.  实际教学进度与计划教学进度相差一周以上、两周以内的；', True, False),
        ('3.  教师酒后进入教学场所（含上课、自习、查寝等），或穿着拖鞋、短裤、背心等不符合教师仪容仪表规范服装上课的；', True,
         False),
        (
        '4.  上课、自习辅导期间接打电话、玩手机（累计5分钟以内）、播放视频时长超过课堂时间三分之一或从事其他与教学无关行为的；',
        True, False),
        ('5.  因教师课堂管理不力，导致课堂秩序混乱的；', True, False),
        ('6.  因教师授课原因，造成授课班级学生出勤率低于80%，或上课期间超过10%的学生玩手机、打瞌睡的；', True, False),
        ('7.  整个学期未布置作业，或作业布置量不足计划量一半的；', True, False),
        ('8.  对学生作业不收、不查、不批阅、不记录的；', True, False),
        ('9.  未按教务科要求及时提交教案、作业、听课记录等常规教学材料的。', True, False),
        ('第四条 考试与监考有下列情形之一的，属于一般教学事故：', False, True),
        ('（一）任课教师或管理人员在考试前向学生泄露试题，造成不良影响的；', True, False),
        ('（二）试题存在错误或试卷准备不足，影响学生正常答题的；', True, False),
        ('（三）监考教师未经批准私自请他人代监考的；', True, False),
        ('（四）监考人员迟到、早退，或工作失误导致试卷分发、回收出现问题的；', True, False),
        ('（五）开考前未按规定清场、检查座位，考试期间未履行监考职责，造成考场秩序混乱的；', True, False),
        ('（六）监考时打瞌睡、聊天或从事其他与监考无关活动的；', True, False),
        ('（七）对考试舞弊行为不予制止，或在考试过程中向学生作出暗示的；', True, False),
        ('（八）不按规定阅卷评分，随意提分、压分或私自更改成绩的；', True, False),
        ('（九）未在规定时间内提交成绩或漏登成绩，影响后续工作的。', True, False),
        ('第五条 管理与保障工作中，有下列情形之一的，属于一般教学事故：', False, True),
        ('（一）因管理人员疏漏导致课程安排冲突，未及时处理而影响教学活动的；', True, False),
        ('（二）未按规定程序擅自变动教学计划的；', True, False),
        ('（三）管理人员未按时开门，影响正常教学活动的；', True, False),
        ('（四）教学设施（灯光、门窗、桌椅、黑板等）已报修但未及时修复，影响教学活动的；', True, False),
        ('（五）多媒体设备故障因管理原因未及时修复的；', True, False),
        ('（六）早晚自习辅导教师管理不严，导致秩序混乱，或擅自离岗的。', True, False),
        ('第六条 课堂教学有下列情形之一的，属于严重教学事故：', False, True),
        ('（一）上课无教案的；', True, False),
        ('（二）教师无故迟到、早退超过5分钟，或整节课未到岗（空堂）的；', True, False),
        ('（三）对学生实施体罚、殴打或使用侮辱性语言，严重伤害学生身心健康的；', True, False),
        ('（四）教学进度严重偏离计划，提前结束新课或滞后两周以上的；', True, False),
        ('（五）因教师失职导致教学过程中发生重大人身伤亡事故的；', True, False),
        ('（六）未经批准擅自停课、调课、请人代课或改为自习达2课时以上，严重影响教学秩序的；', True, False),
        ('（七）因教师授课原因，造成授课班级学生出勤率低于60%，或上课期间超过30%的学生玩手机、打瞌睡的；', True, False),
        ('（八）因课堂歧视后进生，导致学生情绪严重波动或流失的。', True, False),
        ('第七条 考试与监考中有下列情形之一的，属于严重教学事故：', False, True),
        ('（一）泄露试题造成重大影响的；', True, False),
        ('（二）擅自更改成绩或组织学生批阅试卷，造成恶劣影响的；', True, False),
        ('（三）考试结束后未按时上交试卷导致试卷丢失的；', True, False),
        ('（四）未经批准无故不参加监考的；', True, False),
        ('（五）考试结束后未清点试卷，造成试卷丢失且未能追回的。', True, False),
        ('第八条', False, True),
        ('其他人为原因影响教学正常运行的情形，参照本办法第三至第七条予以认定。', False, False),
        ('第九条', False, True),
        (
        '教学中散布违反国家法律法规、教育方针的言论，或传播封建迷信、淫秽内容，造成恶劣影响的，按严重教学事故处理，并依据相关规定给予党纪政纪处分。',
        False, False),
        ('第十条 教学事故的处理程序与措施：', False, True),
        ('（一）教学事故发生后，由教务科会同相关教研室调查核实，约谈当事人后提出初步意见，报教学事故认定领导小组研究决定；',
         True, False),
        ('（二）一般教学事故：当事人须提交书面检查，视情节轻重给予批评教育、内部通报批评，并可处50元罚款；', True, False),
        ('（三）严重教学事故：当事人须提交书面检查，视情节给予全校通报、公开检查、警告直至停职处理，并可处100元罚款；', True,
         False),
        ('（四）教学事故经领导小组认定后，报主管校长批准，出具处罚通知；', True, False),
        ('（五）造成经济损失的，当事人应承担全部或部分赔偿，并可给予相应行政处分；', True, False),
        ('（六）发生严重教学事故者，取消当年评先评优资格。一年内发生两次及以上严重教学事故的，予以辞退；', True, False),
        (
        '（七）一学期内累计发生三次一般教学事故的，按一次严重教学事故处理，并取消当年评先评优资格；累计达到四次一般教学事故的，予以辞退。',
        True, False),
        ('第十一条', False, True),
        ('外聘教师发生教学事故，参照本办法第十条第（二）、（三）款执行；发生严重教学事故的，解除聘任关系。', False, False),
        ('第十二条', False, True),
        ('本办法自发布之日起施行，由学校教务科负责解释。', False, False),
    ]

    # 添加内容到文档
    for text, is_indented, is_bold in content:
        if text.startswith('第') and '条' in text:
            # 条款标题
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = is_bold
            run.font.size = Pt(14)
            run.font.name = '宋体'
        elif is_indented:
            # 缩进段落
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.first_line_indent = Inches(-0.3)
            run = p.add_run(text)
            run.font.size = Pt(12)
            run.font.name = '宋体'
        else:
            # 普通段落
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = is_bold
            run.font.size = Pt(12)
            run.font.name = '宋体'

        # 设置段落间距
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.5

    # 设置页面边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    # 保存文档
    doc.save('xx市实验中等专业学校教学事故认定与处理办法.docx')
    print("文档已生成：xx市实验中等专业学校教学事故认定与处理办法.docx")


if __name__ == "__main__":
    create_formatted_document()